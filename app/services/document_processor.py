import os
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
import aiofiles
import logging
from datetime import datetime, timezone
# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings

from app.core.config import settings
from app.infra.db.chroma_client import chroma_client
from app.infra.repositories.document_repository import DocumentRepository
from app.domain.dtos.document import DocumentStatus, DocumentType

logger = logging.getLogger(__name__)
class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.embeddings = OpenAIEmbeddings(
            openai_api_key=settings.OPENAI_API_KEY,
            model=settings.EMBEDDING_MODEL
        )
    
    async def process_document(
        self, 
        document_id: str, 
        file_path: str, 
        file_type: DocumentType,
        document_repo: DocumentRepository,
        batch_size: int = 128   # embedding batch size
    ):
        """Process uploaded document: extract text, chunk, embed, and store"""
        try:
            await document_repo.update_document_status(document_id, DocumentStatus.PROCESSING)

            # Extract pages with metadata (list of dicts: {"page_number": int, "text": str})
            if file_type == DocumentType.PDF:
                pages = await self._extract_pdf_pages(file_path)
            elif file_type == DocumentType.DOCX:
                txt = await self._extract_docx_text(file_path)
                pages = [{"page_number": 1, "text": txt}]
            elif file_type == DocumentType.TXT:
                txt = await self._extract_txt_text(file_path)
                pages = [{"page_number": 1, "text": txt}]
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Build LangChain-like Documents (keeps metadata per page)
            from langchain.schema import Document as LC_Document
            page_docs = []
            for p in pages:
                md = {
                    "document_id": document_id,
                    "page_number": p["page_number"],
                    "file_name": Path(file_path).name
                }
                page_docs.append(LC_Document(page_content=p["text"], metadata=md))

            # Split into chunks preserving metadata (split_documents keeps metadata)
            split_docs = self.text_splitter.split_documents(page_docs)

            if not split_docs:
                raise ValueError("No content after splitting")

            # Prepare chunk payloads
            chunk_ids = []
            chunk_metadatas = []
            chunk_contents = []
            chunk_docs_for_mongo = []

            now = datetime.now(timezone.utc)
            for i, d in enumerate(split_docs):
                chunk_id = f"{document_id}_chunk_{i}"
                chunk_ids.append(chunk_id)
                chunk_contents.append(d.page_content)
                # merge metadata + chunk-specific metadata
                md = dict(d.metadata or {})
                md.update({
                    "chunk_id": chunk_id,
                    "chunk_index": i,
                    "chunk_size": len(d.page_content),
                    "created_at": now.isoformat()
                })
                chunk_metadatas.append(md)

                chunk_docs_for_mongo.append({
                    "chunk_id": chunk_id,
                    "document_id": document_id,
                    "content": d.page_content,
                    "chunk_index": i,
                    "metadata": md,
                    "created_at": now
                })

            # Persist chunks in MongoDB first (if you prefer to add to chroma first, swap order)
            await document_repo.create_chunks(chunk_docs_for_mongo)

            # Create embeddings in batches to avoid large single requests
            all_embeddings = []
            i = 0
            while i < len(chunk_contents):
                batch = chunk_contents[i:i+batch_size]
                # NOTE: check your embeddings API. If .aembed_documents is not available, call sync method appropriately
                emb = await self._create_embeddings(batch)
                all_embeddings.extend(emb)
                i += batch_size

            # Add to Chroma (use your chroma_client API)
            await chroma_client.add_documents(
                documents=chunk_contents,
                metadatas=chunk_metadatas,
                ids=chunk_ids,
                embeddings=all_embeddings
            )

            # Update status
            await document_repo.update_document_status(
                document_id,
                DocumentStatus.COMPLETED,
                chunks_count=len(chunk_contents)
            )
            logger.info(f"[{document_id}] Processed and stored {len(chunk_contents)} chunks")

        except Exception as e:
            logger.exception(f"[{document_id}] Error processing document: {e}")
            await document_repo.update_document_status(document_id, DocumentStatus.FAILED)
            raise
    
    async def _extract_text(self, file_path: str, file_type: DocumentType) -> str:
        """Extract text from different file types"""
        if file_type == DocumentType.PDF:
            return await self._extract_pdf_text(file_path)
        elif file_type == DocumentType.DOCX:
            return await self._extract_docx_text(file_path)
        elif file_type == DocumentType.TXT:
            return await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_pages(self, file_path: str) -> List[Dict[str, Any]]:
        """Return list of page dicts: [{'page_number': int, 'text': str}, ...]"""
        pages = []
        # read entire file as bytes asynchronously
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()

        import io
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))

        for idx, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            # strip control characters
            text = text.strip()
            pages.append({"page_number": idx, "text": text})

        return pages
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()
    
    async def _create_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Create embeddings for text chunks"""
        try:
            embeddings = await self.embeddings.aembed_documents(texts)
            return embeddings
        except Exception as e:
            logger.error(f"Error creating embeddings: {e}")
            raise

# Global processor instance
document_processor = DocumentProcessor()